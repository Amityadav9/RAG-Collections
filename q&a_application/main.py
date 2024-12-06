import streamlit as st
import faiss
import os
import re
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.llms import Ollama
from streamlit.runtime.uploaded_file_manager import UploadedFile
from dotenv import load_dotenv
import pandas as pd
import time
import logging

# Load environment variables from .env file
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data  # Input is already a text string
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    elif input_type == "CSV":
        if isinstance(input_data, BytesIO):
            df = pd.read_csv(input_data)
        elif isinstance(input_data, UploadedFile):
            df = pd.read_csv(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for CSV")
        documents = []
        for _, row in df.iterrows():
            if row['question'] == row['original_question']:
                clean_answer = row['answer'].strip()
                # Remove multiple spaces but keep single spaces
                clean_answer = ' '.join(clean_answer.split())
            
            qa_pair = f"""Question: {row['question']}\nAnswer: {clean_answer}"""
            documents.append(qa_pair)
    
        documents = "\n\n".join(documents)

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100, separator="\n")
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]  # Access page_content from each Document 
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cuda'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    # Create FAISS vector store with the embedding function
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)  # Add documents to the vector store
    return vector_store

def answer_question(vectorstore, query):
    """Answers a question based on the provided vectorstore."""
    try:
        # Add timeout for Ollama connection
        import requests
        from requests.exceptions import RequestException
        
        # First verify Ollama is accessible
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=20)
            if response.status_code != 200:
                raise ConnectionError("Ollama service is not responding correctly")
        except RequestException as e:
            st.error("Cannot connect to Ollama service. Please ensure Ollama is running.")
            logger.error(f"Ollama connection error: {str(e)}")
            return {
                "query": query,
                "result": "Error: Cannot connect to Ollama service. Please ensure Ollama is running at http://localhost:11434"
            }

        st.write("Initializing Ollama...")
        llm = Ollama(
            model="llama3.2:3b",
            temperature=0.6,
            base_url="http://localhost:11434",
            timeout=120 # Add timeout for Ollama operations
            
        )
        st.write("Ollama initialized successfully")
    
        enhanced_query = (
            f'''Given the question: {query}
    Please provide only the direct answer from the knowledge base without any additional text or explanations.
    If multiple similar answers are found, provide only the most relevant one.'''
        )
        st.write("Enhanced query created.")
    
        qa = RetrievalQA.from_chain_type(
            llm=llm, 
            retriever=vectorstore.as_retriever(),
            return_source_documents=True  # Add this to help with debugging
        )
        st.write("RetrievalQA chain created.")
    
        # Start timing the query processing
        start_time = time.time()
        result = qa({"query": enhanced_query})
        elapsed_time = time.time() - start_time
        st.write(f"Query processed by QA chain in {elapsed_time:.2f} seconds.")
        
        # Clean up the response
        answer = result['result']
        
        # Extract just the answer, removing Question/Answer prefixes
        if 'Answer:' in answer:
            answer = answer.split('Answer:')[-1].strip()
        
        # Remove any remaining Question/Answer patterns
        answer = re.sub(r'Question:.*?Answer:', '', answer, flags=re.DOTALL)
        
        # Clean up any remaining artifacts
        answer = answer.strip()
        answer = ' '.join(answer.split())  # This keeps single spaces between words
        
        st.write("Answer cleaned and ready to display.")
        return {"query": query, "result": answer}

    except Exception as e:
        logger.error(f"Error generating answer: {str(e)}")
        st.error(f"Error generating answer: {str(e)}")
        # Provide more specific error messages based on the exception type
        error_message = str(e)
        if "connection" in error_message.lower():
            error_message = "Connection error: Please ensure Ollama is running and accessible"
        elif "model" in error_message.lower():
            error_message = "Model error: Please ensure the llama3.2:3b model is downloaded in Ollama"
        
        return {
            "query": query,
            "result": f"Error: {error_message}\nPlease try again or contact support if the issue persists."
        }

def main():
    st.title("RAG Q&A App")
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT", "CSV"])
    if input_type == "Link":
        number_input = st.number_input(min_value=1, max_value=20, step=1, label="Enter the number of Links")
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_input("Enter the text")
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])
    elif input_type == 'CSV':
        input_data = st.file_uploader("Upload a CSV file", type=['csv'])
    
    if st.button("Proceed"):
        vectorstore = process_input(input_type, input_data)
        st.session_state["vectorstore"] = vectorstore
    
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            answer = answer_question(st.session_state["vectorstore"], query)
            st.write(f"Answer: {answer['result']}")  # This will show just the answer

if __name__ == "__main__":
    main()